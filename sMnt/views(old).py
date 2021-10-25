#util imports
import datetime
from dateutil.parser import parse
date = datetime.datetime.now().strftime('%Y')
# Imports for main logic
import io
import urllib.parse
from .models import *
from django import forms
from docx import Document
from django.views import View
from django.db.models import Q
from django.shortcuts import render
from django.shortcuts import redirect
from django.contrib.auth.models import User
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from django.contrib.auth import get_user_model
from django.contrib.auth import authenticate, login
from django.http import HttpResponse, FileResponse, JsonResponse
from django.contrib.auth.hashers import make_password, check_password
def sMntSetQuantity(request):
	try:
		if request.method == 'POST':
			init = Init.objects.all().first()
			if not bool(init):
				init = Init(initial_serial_number='0',final_serial_number=request.POST['quant'] )
				init.save()
			else:
				init.final_serial_number = request.POST['quant'] 
				init.save()
			return sMntDashboard().get(request,success=True, error=False, message="Quantity successfully set to " + init.final_serial_number)
		else:
			return redirect('sMntDashboard')
	except Exception as e:
		print('\n\n\n',str(e),'\n\n\n')
		response = HttpResponse()
		response.status_code = 500
		return response
class sMntDashboard(View):
	template = 'sMntDashboard.html'
	def get(self, request, *args, **kwargs):
                try:
                        if request.user.username == 'Admin' or request.user.username == 'IT SUPPORT':
                                statements = Statement.objects.all().order_by('-final_serial_number')[:20]
                                cancelled_statements = Cancelled.objects.all()[:20]
                                user = get_user_model()
                                branches = user.objects.filter(~Q(username='Admin') & ~Q(username='beod') & ~Q(username='IT SUPPORT'))
                                if bool(kwargs.get('error')):
                                        error = 'true'
                                        success = 'false'
                                        message = kwargs.get('message')
                                elif bool(kwargs.get('success')):
                                        error="false"
                                        success="true"
                                        message = kwargs.get('message')
                                else:
                                        error = 'false'
                                        success='false'
                                        message = ''
                                init = Init.objects.all().first()
                                #Incase the init table is empty for e.g. when manually reinitializing system.
                                if not bool(init):
                                        init = Init(initial_serial_number='0',final_serial_number='0',current_serial_number='0')
                                        init.save()
                                return render(request,self.template,{'date':date,'branches':branches, 'Init':init,'error':error,'message':message, 'success':success})
                        else:
                                return redirect('redirectLandingPage')
                except Exception as e:
                        print('\n\n\n',str(e),'\n\n\n')
                        response = HttRepsonse()
                        response.status_code = 500 
                        return response
	def post(self, request, *args, **kwargs):
		print('POST request made to sMntDashboard.')
		return redirect('sMntDashboard')
def sMntAddRecord(request):
	print('\n\n\n', request.POST, '\n\n\n\n')
	try:
		form = urllib.parse.parse_qs( request.POST['statement'] )
		if ',' in form['initial_serial_number'][0]:
			form['initial_serial_number'][0] = form['initial_serial_number'][0].replace(',','')
		if ',' in form['final_serial_number'][0]:
			form['final_serial_number'][0] = form['final_serial_number'][0].replace(',','')
		init = Init.objects.all()[0]
		final_serial_number = '' 
		current_serial_number = ''
		submitted_final_serial_number = int(form['final_serial_number'][0])
		submitted_initial_serial_number = int(form['initial_serial_number'][0])
		current_serial_number = int(init.initial_serial_number)
		final_serial_number = int(init.final_serial_number)
		print(current_serial_number,submitted_initial_serial_number, final_serial_number, submitted_final_serial_number)
		if submitted_final_serial_number >= final_serial_number or submitted_initial_serial_number >= final_serial_number:
			return JsonResponse({'error':True,'message':'The serial numbers entered are greater than the maximum quantity.'})
		elif  submitted_final_serial_number < current_serial_number or submitted_initial_serial_number < current_serial_number:
			return JsonResponse({'error':True,'message':'The serial numbers enetered are below the current serial no.'})
		statement = Statement(
			branch=form['branch'][0],
			customer_name=form['customer_name'][0],
			acc_no=form['account_no'][0],
			initial_serial_number=form['initial_serial_number'][0],
			final_serial_number=form['final_serial_number'][0],
			date=datetime.datetime.now(),
		)
		statement.save()
		init = Init.objects.all().first()
		init.initial_serial_number = statement.final_serial_number 
		init.save()
		quant = Quant.objects.filter(date=datetime.datetime.now()).first()
		if bool(quant):
			quant.current_quant = init.final_serial_number
			quant.total_usage = statement.final_serial_number
			quant.save()
		else:
			quant = Quant(current_quant=init.final_serial_number,total_usage=statement.final_serial_number, date=datetime.datetime.now())
			quant.save()
		return JsonResponse({
			'date':statement.date.strftime('%a %d %b %Y'),
			'branch':statement.branch,
			'customer_name':statement.customer_name,
			'account_no':statement.acc_no,
			'initial_serial_number':statement.initial_serial_number,
			'final_serial_number':statement.final_serial_number,
			'usage':str( (int(statement.final_serial_number) - int(statement.initial_serial_number) ) + 1 ),
			'remaining_stock':( int(init.final_serial_number) - int (init.initial_serial_number) ) + 1,
			'total_usage':init.initial_serial_number,
			'current_serial_number':str(int(init.initial_serial_number) + 1),
		})
	except Exception as e:
		print('\n\n\n',str(e),'\n\n\n')
		return JsonResponse({'error':True,'message':'Could not add record.'})
def sMntCancelStatments(request):
	try:
                form = urllib.parse.parse_qs( request.POST['statement'] )
                if ',' in form['Cfinal_serial_number'][0]:
                        form['Cfinal_serial_number'][0] = form['Cfinal_serial_number'][0].replace(',','')
                if ',' in form['Cinitial_serial_number'][0]:
                        form['Cinitial_serial_number'][0] = form['Cinitial_serial_number'][0].replace(',','')
                init = Init.objects.all().first()
                if not bool(init):
                        return JsonResponse({'error':True,'message':'The quantity has not yet been set!!'})
                final_serial_number = ''
                initial_serial_number = ''
                submitted_initial_serial_number = int(form['Cinitial_serial_number'][0])
                submitted_final_serial_number = int(form['Cfinal_serial_number'][0])
                initial_serial_number = int(init.initial_serial_number)
                final_serial_number = int(init.final_serial_number) 
                if submitted_initial_serial_number >= final_serial_number or submitted_final_serial_number >= final_serial_number:
                        return JsonResponse({'error':True,'message':'The serial numbers entered exceed the maximum quantity.'})
                if submitted_final_serial_number < initial_serial_number + 1 or submitted_initial_serial_number < initial_serial_number + 1:
                        return JsonResponse({'error':True, 'message':'The serial numbers entered are below the current serial no.'})
                cancelled_statement = Cancelled(
                        initial_serial_number=form['Cinitial_serial_number'][0],
                        final_serial_number=form['Cfinal_serial_number'][0],
                        cancellation_date=datetime.datetime.now()
                )
                cancelled_statement.save()
                init = Init.objects.all().first()
                init.initial_serial_number = cancelled_statement.final_serial_number 
                init.save()
                quant = Quant.objects.filter(date=datetime.datetime.now()).first()
                if bool(quant):
                        quant.current_quant = init.final_serial_number
                        quant.total_usage = cancelled_statement.final_serial_number
                        quant.save()
                else:
                        quant = Quant(current_quant=init.final_serial_number, total_usage=cancelled_statement.final_serial_number, date=datetime.datetime.now())
                        quant.save()
                return JsonResponse({
                        'initial_serial_number':cancelled_statement.initial_serial_number,
                        'final_serial_number':cancelled_statement.final_serial_number,
                        'useage':str( int(cancelled_statement.final_serial_number) - int (cancelled_statement.initial_serial_number) + 1),
                        'remaining_stock':( int(init.final_serial_number) - int (init.initial_serial_number) ) + 1,
                        'total_usage':init.initial_serial_number,
                        'current_serial_number':str(int(init.initial_serial_number) + 1),
                })
	except Exception as e:
		print('\n\n\n',str(e),'\n\n\n')
		return JsonResponse({'errror':True,'message':'Could not cancel statement.'})
def sMntGenReport(request):
	try:
		initial_date = datetime.datetime.strptime(request.POST['initialDate'], '%Y-%m-%d')
		initial_date = initial_date.replace(tzinfo=datetime.timezone.utc)
		if request.POST['for'] == 'customer_statements':
			if request.POST['branch'] == 'all_branches':
				customer_statements = Statement.objects.filter(date__date=initial_date)
			else:  
				customer_statements = Statement.objects.filter(date__date=initial_date, branch=request.POST['branch'])
			if bool(customer_statements):
				todays_usage = 0
				total_cancelled = 0
				total_printed = 0
				total_usage = 0
				lowest_serial_number = None
				document = Document()
				document.add_heading('REPORT FOR STATEMENTS PRINTED \n' + str( datetime.datetime.strptime(request.POST['initialDate'],'%Y-%m-%d').strftime('%d-%m-%Y') ), 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
				table = document.add_table(1,6)
				table.autofit = True
				table.alignment = WD_TABLE_ALIGNMENT.CENTER
				table.cell(0,0).add_paragraph().add_run('Branch').bold = True
				table.cell(0,1).add_paragraph().add_run('Customer Name').bold = True
				table.cell(0,2).add_paragraph().add_run('Account Number').bold = True
				table.cell(0,3).add_paragraph().add_run('Initial Serial Number').bold = True
				table.cell(0,4).add_paragraph().add_run('Final Serial Number').bold = True
				table.cell(0,5).add_paragraph().add_run('Usage').bold = True
				for statement in customer_statements:
					if total_usage < int(statement.final_serial_number):
						total_usage = int(statement.final_serial_number)
					if lowest_serial_number == None or lowest_serial_number > int(statement.initial_serial_number):
						lowest_serial_number = int(statement.initial_serial_number)
					useage = str(int(statement.final_serial_number) - int(statement.initial_serial_number) + 1)
					total_printed = total_printed + int(useage)
					statement_details = [statement.branch,statement.customer_name,statement.acc_no,statement.initial_serial_number,statement.final_serial_number,useage]
					table.add_row()
					row_index = len(table.rows)
					for index,column in enumerate(statement_details): 
						table.cell(row_index-1,index).add_paragraph(statement_details[index])
				todays_usage = total_printed
				for statement in Cancelled.objects.filter(cancellation_date__date=initial_date):
					total_cancelled = total_cancelled + (int(statement.final_serial_number) - int(statement.initial_serial_number) + 1)
					if total_usage < int(statement.final_serial_number):
						total_usage = int(statement.final_serial_number)
					if lowest_serial_number == None or lowest_serial_number > int(statement.initial_serial_number):
						lowest_serial_number = int(statement.initial_serial_number)
				todays_usage = todays_usage + total_cancelled
				document.add_page_break()
				document.add_heading('Summary',level=1)
				summary_table = document.add_table(3,6)
				summary_table.autofit = True
				summary_table.cell(0,0).add_paragraph().add_run('Current Serial No').bold = True
				summary_table.cell(0,1).add_paragraph().add_run('Quantity').bold = True
				summary_table.cell(0,2).add_paragraph().add_run('Total Usage').bold = True
				summary_table.cell(0,3).add_paragraph().add_run('Remaining Stock').bold = True
				summary_table.cell(0,4).add_paragraph().add_run('Today\'s Usage').bold = True
				summary_table.cell(0,5).add_paragraph().add_run('Total Printed').bold = True
				init = Init.objects.all()[0]
				quantity_at_that_date = Quant.objects.get(date=initial_date)
				summary_table.cell(1,0).add_paragraph(str(total_usage+1))
				summary_table.cell(1,1).add_paragraph(quantity_at_that_date.current_quant)
				summary_table.cell(1,2).add_paragraph(str(total_usage))
				summary_table.cell(1,3).add_paragraph(str( int(quantity_at_that_date.current_quant) - total_usage + 1 ) )
				summary_table.cell(1,4).add_paragraph(str(todays_usage))
				summary_table.cell(1,5).add_paragraph( str(total_printed) )
				document.add_heading('I.T MANAGER\t\t\t\t\t\t\tACCOUNTANT', level=2)
				f = io.BytesIO()
				document.save(f)
				f.seek(0)
				return FileResponse(f, as_attachment=True, filename='StatementsReport.docx')
			else:
				return sMntDashboard().get(request,error=True,success=False,message='There are no matching statement entries for the selected date.')
		elif request.POST['for'] == 'cancelled_statements':
			document = Document()
			document.add_heading('REPORT FOR STATEMENTS CANCELLED \n' + str( datetime.datetime.strptime(request.POST['initialDate'],'%Y-%m-%d').strftime('%d-%m-%Y') ), 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
			cancelled_statements = Cancelled.objects.filter(cancellation_date__date=initial_date)
			if bool(cancelled_statements):
				if not bool(cancelled_statements):
					document.add_heading('Cancelled Statements', level=1)
					document.add_paragraph('There are no cancelled statements.')
				else:
					total_cancelled = 0
					total_usage = 0
					lowest_serial_number = None
					document.add_heading('Cancelled Statements', level=1)
					cancelled_statements_table = document.add_table(1,3)
					cancelled_statements_table.autofit = True
					cancelled_statements_table.alignment = WD_TABLE_ALIGNMENT.CENTER
					cancelled_statements_table.cell(0,0).add_paragraph().add_run('Initial Serial Number').bold = True
					cancelled_statements_table.cell(0,1).add_paragraph().add_run('Final Serial Number').bold = True
					cancelled_statements_table.cell(0,2).add_paragraph().add_run('Usage').bold = True
					for statement in cancelled_statements:
						if total_usage < int(statement.final_serial_number):
							total_usage = int(statement.final_serial_number)
						if lowest_serial_number == None or lowest_serial_number > int(statement.initial_serial_number):
							lowest_serial_number = int(statement.initial_serial_number)
						useage = str(int(statement.final_serial_number) - int(statement.initial_serial_number) + 1)
						total_cancelled = total_cancelled + int(useage)
						statement_details = [statement.initial_serial_number,statement.final_serial_number,useage]
						cancelled_statements_table.add_row()
						row_index = len(cancelled_statements_table.rows)
						for index,column in enumerate(statement_details): 
							cancelled_statements_table.cell(row_index-1,index).add_paragraph(statement_details[index])
					for statement in Statement.objects.filter(date__date=initial_date):
						if total_usage < int(statement.final_serial_number):
							total_usage = int(statement.final_serial_number)
						if lowest_serial_number == None or lowest_serial_number > int(statement.initial_serial_number):
							lowest_serial_number = int(statement.initial_serial_number)
					document.add_page_break()
					document.add_heading('Summary',level=1)
					summary_table = document.add_table(3,6)
					summary_table.autofit = True
					summary_table.cell(0,0).add_paragraph().add_run('Current Serial No').bold = True
					summary_table.cell(0,1).add_paragraph().add_run('Quantity').bold = True
					summary_table.cell(0,2).add_paragraph().add_run('Total Usage').bold = True
					summary_table.cell(0,3).add_paragraph().add_run('Remaining Stock').bold = True
					summary_table.cell(0,4).add_paragraph().add_run('Today\'s Usage').bold = True
					summary_table.cell(0,5).add_paragraph().add_run('Total Cancelled').bold = True
					init = Init.objects.all()[0]
					quantity_at_that_date = Quant.objects.get(date=initial_date)
					summary_table.cell(1,0).add_paragraph(str(total_usage+1))
					summary_table.cell(1,1).add_paragraph(quantity_at_that_date.current_quant)
					summary_table.cell(1,2).add_paragraph(str(total_usage))
					summary_table.cell(1,3).add_paragraph(str( int(quantity_at_that_date.current_quant) - total_usage + 1 ) )
					summary_table.cell(1,4).add_paragraph(str((int(quantity_at_that_date.total_usage)-lowest_serial_number)+1))
					summary_table.cell(1,5).add_paragraph( str(total_cancelled) )
					document.add_heading('I.T MANAGER\t\t\t\t\t\t\tACCOUNTANT', level=2)
					f = io.BytesIO()
					document.save(f)
					f.seek(0)
					return FileResponse(f, as_attachment=True, filename='StatementsReport.docx')
			else:
				return sMntDashboard().get(request,error=True,success=False,message='There are no matching cancelled statements for the selected date.')
		else:
			response = HttpResponse()
			response.status_code = 503
			return response
	except Exception as e:
		print('\n\n\n',str(e),'\n\n\n')
		return redirect('sMntDashboard')
def sMntSearch(request):
	try:
		if request.method == 'POST':
			if request.POST['for'] == 'customer_statements':
				query = request.POST['search[value]']
				query = Q(branch__contains=query) | Q(customer_name__icontains=query) | Q(acc_no__contains=query) | Q(initial_serial_number__contains=query) | Q(final_serial_number__contains=query) | Q(date__date__contains=query)
				holder = Statement.objects.filter(query).order_by('-final_serial_number')
				statements = holder[int(request.POST['start']):int(request.POST['start']) + int(request.POST['length'])]
				if bool(statements):
					rows = [[statement.date.strftime('%A %d. %B %Y'),statement.branch,statement.customer_name,statement.acc_no,statement.initial_serial_number,statement.final_serial_number,str((int(statement.final_serial_number)-int(statement.initial_serial_number))+1)]for statement in statements]
					return JsonResponse({'draw':int(request.POST['draw']),'recordsTotal':Statement.objects.all().count(), 'recordsFiltered':holder.count(),'data':rows})
				else:
					return JsonResponse({'recordsTotal':Statement.objects.all().count(), 'recordsFiltered':0,'data':[]})
			if request.POST['for'] == 'cancelled_statements':
				query = request.POST['search[value]']
				query = Q(initial_serial_number__contains=query) | Q(final_serial_number__contains=query) | Q(cancellation_date__date__contains=query)
				holder = Cancelled.objects.filter(query).order_by('-final_serial_number')
				cancelled_statements = holder[int(request.POST['start']):int(request.POST['start']) + int(request.POST['length'])]
				if bool(cancelled_statements):
					rows = [[cancelled_statement.cancellation_date.strftime('%A %d. %B %Y'),cancelled_statement.initial_serial_number,cancelled_statement.final_serial_number,str((int(cancelled_statement.final_serial_number)-int(cancelled_statement.initial_serial_number))+1)]for cancelled_statement in cancelled_statements]
					return JsonResponse({'draw':int(request.POST['draw']),'recordsTotal':Cancelled.objects.all().count(), 'recordsFiltered': holder.count(),'data':rows})
				else:
					return JsonResponse({'recordsTotal':Cancelled.objects.all().count(), 'recordsFiltered':0,'data':[]})
		else:
			return redirect('redirectLandingPage')
	except Exception as e:
		print('\n\n\n',str(e),'\n\n\n')
		return redirect('redirectLandingPage')
class sMntDelSmnt(View):
        def get(self, request, *args, **kwargs):
                return redirect('sMntDashboard')
        def post(self, request, *args, **kwargs):
                try:
                    init = Init.objects.all().first()
                    previous_highest_serial_number = str(int(init.initial_serial_number))
                    statement = Statement.objects.filter(final_serial_number = previous_highest_serial_number ).first()
                    cancelled = Cancelled.objects.filter(final_serial_number = previous_highest_serial_number).first()
                    if bool(statement):
                            init.initial_serial_number = str(int(statement.initial_serial_number) - 1)
                            statement.delete()
                            init.save()
                            return sMntDashboard().get(request, error=False, success=True, message='Statement for ' + statement.customer_name + ' with initial serial number: ' + statement.initial_serial_number + ' and final serial number: ' + statement.final_serial_number  + ' has been removed.'  )
                    elif bool(cancelled):
                            init.initial_serial_number = str(int(cancelled.initial_serial_number) - 1)
                            cancelled.delete()
                            init.save()
                            return sMntDashboard().get(request, error=False, success=True, message='Cancelled Statement with initial serial number: ' + cancelled.initial_serial_number + ' and final serial number ' + cancelled.final_serial_number  + ' has been removed.'  )
                    raise Exception()
                except Exception as e:
                    print('\n\n\n',str(e),'\n\n\n')
                    return sMntDashboard().get(request, error=True, success=True, message='Could not delete record.' )
# for val in Statement.objects.all():
# 	if val.customer_name == 'SALAMATU KARGBO':
# 		val.delete()
# for val in Cancelled.objects.all():
# 	if val.initial_serial_number == '8777777778':
# 		val.delete()
# for val in Init.objects.all():
# 	val.initial_serial_number = '203732'
# 	val.save()
# for val in Quant.objects.all()[1:]:
# 	val.delete()
# quant = Quant(date=datetime.datetime.strptime('02-23-2021','%m-%d-%Y'), current_quant='204000',total_usage = '203732')
# quant.save()
# quant = Quant.objects.filter(date__date=datetime.datetime.strptime('23-02-2021','%d-%m-%Y'))
# quant[0].delete()
# quant[1].delete()
# quant[0].total_usage = '203732'
# quant[0].save()
# init = Init.objects.all()[0]
# init.initial_serial_number = '203732'
# init.final_serial_number = '204000'
# init.current_serial_number = '203733'
# init.save()
# for val in Quant.objects.all():
# 	print(val.date,val.total_usage,val.current_quant)
# for val in Cancelled.objects.all():
# 	for v in Cancelled.objects.all():
# 		check = 0
# 		if val is val.final_serial_number == v.final_serial_number or val.initial_serial_number == v.initial_serial_number :
# 			check = check + 1
# 		if check > 1:
# 			v.delete() 