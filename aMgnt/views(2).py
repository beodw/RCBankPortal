######################## Utility methods ######################
import datetime
from dateutil.parser import parse

date = datetime.datetime.now().strftime('%Y')

def is_date(string, fuzzy=False):
    try: 
        parse(string, fuzzy=fuzzy)
        return True

    except ValueError:
        return False

#Clean form. Easier than implementing django clean() method in forms.py for each form.
def unpack(submitted_form):
	submitted_form = dict(submitted_form) #Gen new dict because request.POST is immutable
	submitted_form.pop('csrfmiddlewaretoken')
	# ensure field values are single string and not a list.
	for key, val in submitted_form.items():
		if type(val) is not str:
			submitted_form[key] = submitted_form[key][0]
	return submitted_form 

#################  Main Logic Classes and function ####################
import socket
import inflect
from .forms import *
from .models import *
from django import forms
from cherrypy import log
from docx import Document
import dateutil.relativedelta
from django.views import View
from docx.shared import Inches
from django.db.models import Q
from django.shortcuts import render
from django.utils.timezone import make_aware
from django.http import HttpResponse, JsonResponse
from django.shortcuts import redirect
from docx.enum.section import WD_ORIENT
from django.contrib.auth.models import User
from django.core.files.images import ImageFile
from django.contrib.auth import authenticate, login
from django.contrib.auth.hashers import make_password, check_password
#############For genreating pdf reports#########
import io
from reportlab.pdfgen import canvas
from django.http import FileResponse
#############################\
def delete_hardware(request):
	if request.method == "GET":
		response = HttpResponse()
		response.status_code = 400 
		return response
	elif request.method == 'POST':
		try:
			if authenticate(username=request.user.username,password=request.POST['adminPassword']) is None:
				return View_asset().get(request,id=request.POST['id'],type=request.POST['type'],error="true",success="false",message="Wrong admin password")
			else:
				hardware = Hardware.objects.get(id=request.POST['id'])
				if hardware.hardware_type == "SystemUnit":
					hardware = System_unit.objects.get(id=hardware.id)
				elif hardware.hardware_type == "Server":
					hardware = Software_capable_hardware.objects.get(id=hardware.id)
				hardware.delete()
				return View_asset().get(request,id=request.POST['id'],type=request.POST['type'],error="false",success="true",message="Hardware deleted successfully")
		except:
			response = HttpResponse()
			response.status_code = 400 
			return response
def exists(Model, query_list={}):
	qs = Model.objects.filter(**query_list)
	if bool ( qs ):
		return True
	else:
		return False
#only admin has access to this view
class Manage_users(View):
	template = 'manage_users.html'
	def get(self,request, *args, **kwargs):
		if request.user.username == 'Admin':
			from django.contrib.auth import get_user_model
			users = Branch.objects.all()

			context = {'branch_name':request.user.username, 'users':users, 'date':date}
			if bool ( kwargs.get('error') ) or bool( kwargs.get('success') ):
				context.update(kwargs)
			else:
				context.update({'message':'','error':'false', 'success':'false'})

			context['add_branch_form'] = Add_branch_form() 
			context['add_branch_form'].fields['verify_password'].label = "Re-enter Password"
			context['add_branch_form'].fields['admin_password'].label = "Enter admin password"
			return render(request, self.template, context)
		else:
			return redirect('/asset/viewAllAssets/')
	def post(self, request, *args, **kwargs):
		try:
			flag = False
			#check if admin raises error if not admin or admin_password field missing from form.
			if authenticate(username=request.user.username, password=request.POST['admin_password']) is None:
				flag = True
				raise Exception('Incorrect admin password.') 
			details = dict ( request.POST )
			if request.POST['purpose'] == 'add_branch':
				details = unpack(details)
				if details['branch_password'] == details['verify_password'] and not bool(Branch.objects.filter(name=details['branch_name'])):
					password = make_password(details['branch_password'])
					branch = Branch(name=details['branch_name'], password=password)
					branch.save()
					User.objects.create_user(username=details['branch_name'], password=details['branch_password'])
					return self.get(request, **{'success':'true', 'error':'false', 'message':'Branch ' + branch.name + ' successfully added.'})
				else:
					# notify admin user that branch already exists or branch passwords do not match.
					if bool(Branch.objects.filter(name=details['branch_name'])):
						flag = True 
						raise Exception('Branch ' + details['branch_name'] + ' Already Exists.')
					elif not details['branch_password'] == details['verify_password']:
						flag = True 
						raise Exception('Passwords do not match.')
			elif request.POST['purpose'] == 'delete_branch':
			
				for name in details['branches_to_delete']:
					branch = Branch.objects.get(name=name)
					branch.delete()
					user = User.objects.get(username=name)
					user.delete()
					return self.get(request, **{'message':'Branch '+branch.name+' has been deactivated.', 'error':'true', 'success':'false'})
			
			elif request.POST['purpose'] == 'reset_password':
				#Reset password for all branches chosen in table in manage branches template.
				branches = ''
				if details.get('branches_to_reset') is None or len(details['branches_to_reset']) < 1 :#if branches to select is removed fro form submission
					raise Exception('Branches to reset cannot be empty')
				else:
                                    
					for branch in details['branches_to_reset']:
						branch = Branch.objects.get(name=branch)#Check if branch already exists will throw error if not or if more than one in db.
						if branch.is_active is True:
							password = make_password(details['new_branch_password'][0])#because new_branch_password is list
							branch.password = password
							branch.save()
							user = User.objects.get(username=branch.name)
							user.password = password
							user.save()
							branches += branch.name + ', '   
					if branches == '':#if all branches to reset are deactivated branches.
						flag = True
						raise Exception('The selected branch user(s) are deactivated. Their passwords cannot be reset.')
					elif len(details['branches_to_reset']) > 1:
						to_be_conjugation = 'have'
						plural = 'passwords'
					elif len(details['branches_to_reset']) == 1:
						to_be_conjugation = 'has'
						plural = 'password'

					return self.get(request, **{'success':'true', 'error':'false', 'message':'Branch user(s): ' + branches + plural+ ' ' + to_be_conjugation + ' been successfully reset.'})

                    
			return redirect('/asset/manageUsers/')
		except Exception as e:
			exception_message = str(e)
			if exception_message != '' and flag is True:
				return self.get(request, **{'error':'true','success':'false', 'message':exception_message})
			else:
				return self.get(request)

def asset_login(request):
	try:
		if request.method == 'GET': #serve login page
			if request.user.is_authenticated : #except if already logged_in
				return redirect('/asset/viewAllAssets/')
			else :
				context = {'message':'', 'error':'none', 'date':date}
				return render(request, "index.html", context)
		if request.method == 'POST': #validate user
			# if request.POST.get('username') == 'Admin':
			user = authenticate(username=request.POST.get('username'), password=request.POST.get('password') )
			# else:
			# 	user = Branch.objects.get(name=request.POST['username'])
			# 	if check_password(request.POST['password'],user.password):
			# 		login(request, user, backend=user.backend)
			# 		return redirect('asset/viewAllAssets/')
			# 	else:
			# 		context = {'message' : 'Incorrect Username or Password', 'error':'visible', 'date':date}
			# 		return render(request, 'index.html',  context)

			if user is not None :
				login(request, user, backend=user.backend)
				requested_page = request.GET.get('next')
				if requested_page is not None :
					return redirect(requested_page)
				else :
					return redirect('/asset/viewAllAssets/')
			else:
				context = {'message' : 'Incorrect Username or Password', 'error':'visible', 'date':date}
				return render(request, 'index.html',  context)
	except:
		return redirect('/asset/login/')

def asset_logout(request):
	from django.contrib.auth import logout
	logout(request)
	return redirect('landingPageLogin')

class View_asset(View):
	template = 'view_asset.html'
	error_message = 'view_all_assets.html'
	error = 0

	def init_context(self, error=0, message=''):
		context = {'date':'date', 'error':error}
		return context

	def resolve_incident(self, details):
		incident = Incident.objects.get(id=details['incident_id'])

		incident.steps_taken_to_resolve = details['steps_to_resolve_description']
		incident.date_of_resolution = datetime.datetime.now().strftime('%Y-%m-%d') #defined above
		incident.status = 'Resolved'

		incident.save()

		# Check if there are any unresolved severe incidents
		if bool ( Incident.objects.filter(hardware=incident.hardware.id, is_severe=1, status='Not Resolved' ) ):
			incident.hardware.condition = 'Bad'
		# Check if there are any non_severe unresolved incidents
		elif bool ( Incident.objects.filter(hardware=incident.hardware.id, is_severe=0, status='Not Resolved' ) ):
			incident.hardware.condition = 'Mid'
		else:
			incident.hardware.condition = 'Good'

		incident.hardware.save()

		return redirect('/asset/viewAsset/' + str ( details['id']  ) + '/' + str ( details['hardware_type'] ) )

	def new_incident(self, details):
		details.pop('purpose')
		if details['hardware_type'] == 'Server':
			hardware = Software_capable_hardware.objects.get(id=details['id'])
		elif details['hardware_type'] == 'SystemUnit':
			hardware = System_unit.objects.get(id=details['id'])
		else:
			hardware = Hardware.objects.get(id=details['id'])

		details.pop('hardware_type')
		details.pop('id')#so incident doesn't take the same id as hardware.
		Incident(hardware=hardware, date_of_resolution=None, **details).save()
		#Update hardware condition
		if details['is_severe'] == 'True':
			hardware.condition = 'Bad'
	
		elif hardware.condition == 'Good' and details['is_severe'] == 'False':
			hardware.condition = 'Mid'
		hardware.save()

		return redirect('/asset/viewAsset/'+ str ( hardware.id  )+ '/' + str ( hardware.hardware_type) )

	#Update hardware based on hardware type of post request.
	def update_hardware(self, details):
		details.pop('purpose')
		#Update details by making a model obj and saving with same id as already existing obj.
		if details['hardware_type'] == 'Server':
			if Software_capable_hardware_form_components(details).is_valid():
				hardware = Software_capable_hardware(**details)
		elif details['hardware_type'] == 'SystemUnit':
			if System_unit_form_components(details).is_valid():
				hardware = System_unit(**details)
		else:
				hardware = Hardware(**details)

		hardware.supply = Supply(id=details['supply_id'])
		hardware.is_received = True
		hardware.save()

		return redirect('/asset/viewAsset/' + str( hardware.id ) + '/' + hardware.hardware_type)

	def check_if_replacement_exists(self, hardware):
		try:
			return Replacement.objects.get(new_hardware_id=hardware.id)
		except:
			return None

	def gen_replacement_history(self, hardware, list_of_replacements):
		replacement = self.check_if_replacement_exists(hardware)
		if not bool( replacement ):
			return list_of_replacements
		else:
			list_of_replacements.append( replacement )

			return self.gen_replacement_history(hardware=replacement.old_hardware , list_of_replacements=list_of_replacements)
		
	def replace_hardware(self, request, details):
		# Set is_replaced field of old hardware to true
		# Create new hardware
		# create new replacement using old hawrdware and new hardware details and the current date
		details.pop('purpose')
		if System_unit_form_components(details).is_valid():
			old_hardware = System_unit.objects.get( id=details['id'] )
			old_hardware.is_replaced = True

			details.pop('id')
			new_hardware = System_unit(**details)

		elif Software_capable_hardware_form_components(details).is_valid():
			old_hardware = Software_capable_hardware.objects.get( id=details['id'] )
			old_hardware.is_replaced = True
			
			details.pop('id')#id of old hardware
			new_hardware = Software_capable_hardware(**details)

		else:
			old_hardware = Hardware.objects.get( id=details['id'] )
			old_hardware.is_replaced = True

			details.pop('id')
			new_hardware = Hardware(**details)
		try:
			# New hardware must be saved first beacuse it may have a serial number that already exists depending on user input. We need to catch that.
			new_hardware.supply = old_hardware.supply
			new_hardware.is_received = True
			new_hardware.save()
			old_hardware.save()

			Replacement(old_hardware=old_hardware, new_hardware=new_hardware, date_of_replacement=datetime.datetime.now().strftime('%Y-%m-%d')).save()

			return redirect('/asset/viewAsset/' + str ( new_hardware.id  ) + '/' + str ( new_hardware.hardware_type) )
		except Exception as e:
			#print(e, '\n\n\n\n\n')
			error=1
			message='A Hardware With That Serial Number Already Exists'
			return self.get(request , **{'id':old_hardware.id, 'type': old_hardware.hardware_type, 'error':error, 'message':message})

	def generic_update_form(self, hardware):
		generic_hardware_form = Update_generic_hardware_form(initial=hardware)
		# generic_hardware_form.fields['assigned_to'].widget.attrs['readonly'] = ''
		# generic_hardware_form.fields['branch'].widget.attrs['readonly'] = ''
		# generic_hardware_form.fields['date_of_supply'].widget.attrs['readonly'] = ''
		generic_hardware_form.fields['hardware_type'].widget.attrs['readonly'] = ''
		generic_hardware_form.fields['serial_number'].widget.attrs['readonly'] = ''
		generic_hardware_form.fields['brand_and_model'].widget.attrs['readonly'] = ''
		generic_hardware_form.fields['hardware_type'].label = ''
		# generic_hardware_form.fields.pop('department')

		return generic_hardware_form

	def get(self, request, *args, **kwargs):
		context = {'date':date, 'error':self.error, 'message':self.error_message}
		context['branch_name'] = request.user.username
		context['hardware_id'] = kwargs['id']
		try:
			if kwargs['type'] == 'Server':
				hardware = Software_capable_hardware.objects.get(id=kwargs['id'], is_received=True, is_replaced=False)
				holder = hardware.__dict__
				# Generate replacement history
				context['list_of_replacements'] = self.gen_replacement_history(hardware, [])
				# hardware = hardware.__dict__
				holder.pop('hardware_ptr_id')
				#Update form components for Server and other software capable hardware.
				context['software_capable_hardware_form_components'] = Software_capable_hardware_form_components(initial=holder)
				#Form components for replace hardware modal
				context['add_software_capable_hardware_components'] = Software_capable_hardware_form_components()
				context['hardware_type'] = kwargs['type']
			elif kwargs['type'] == 'SystemUnit' or kwargs['type'] == 'Laptop':
				hardware = System_unit.objects.get(id=kwargs['id'], is_received=True, is_replaced=False)
				holder = hardware.__dict__
				if kwargs['type'] == "Laptop":
					#holder.pop('keyboard_serial_number')
					#holder.pop('monitor_serial_number')
					holder.pop('service_tag_code')
				# Generate replacement history
				context['list_of_replacements'] = self.gen_replacement_history(hardware, [])
				# hardware = hardware.__dict__
				holder.pop('hardware_ptr_id')
				holder.pop('software_capable_hardware_ptr_id')
				#Update form components for system unit.
				context['software_capable_hardware_form_components'] = Software_capable_hardware_form_components(initial=holder)
				context['update_system_unit_hardware_form_components'] = System_unit_form_components(initial=holder)
				#Form components for replace hardware modal
				context['add_software_capable_hardware_components'] = Software_capable_hardware_form_components() #becasuse Software_capable_hardware_form_components declared above already has prepopulated fields.
				context['system_unit_hardware_form_components'] = System_unit_form_components()

				context['hardware_type'] = kwargs['type']
			else:
				hardware = Hardware.objects.get(id=kwargs['id'], is_received=True, is_replaced=False)
				holder = hardware.__dict__

				# Generate replacement history
				context['list_of_replacements'] = self.gen_replacement_history(hardware, [])

				# hardware = hardware.__dict__ 

			context['incidents'] = Incident.objects.filter(hardware=hardware)

			# Raise exception if user is trying to access a replaced hardware directly using url.
			if hardware.is_replaced is True:
				raise Exception('Hardware is replaced is True')
		#Handle invalid get request or errors in connecting to db.
		except Exception as e:
			print('\n\n\n\n\n\n\n' , e, '\n To Handle invalid get request or errors in connecting to db.')
			return render(request,'generic404page.html',{'date':date})
        #Generic update form components
		context['update_generic_hardware_form'] = self.generic_update_form(holder)

		#Generic replace hardware form components.
		context['replace_hardware_generic_form'] = Replace_hardware_form(initial={ 'hardware_type':hardware.hardware_type, })
		context['replace_hardware_generic_form'].fields['hardware_type'].widget.attrs['hidden'] = '' 
		context['replace_hardware_generic_form'].fields['hardware_type'].widget.attrs['id'] = 'replace_hardware_input_hardware_type'
		context['replace_hardware_generic_form'].fields['brand_and_model'].widget.attrs['id'] = 'replace_hardware_input_brand_and_model'
		context['replace_hardware_generic_form'].fields['hardware_type'].label = ''

		context['hardware'] = holder
		context['hardware_branch'] = hardware.supply.branch
		context['hardware_type'] = kwargs['type']

		if bool ( kwargs.get('error') ) :
			context['error'] = kwargs['error']
			context['success'] = kwargs['success']
			context['message'] = kwargs['message']
		elif bool ( kwargs.get('success') ):
			context['success'] = kwargs['success']
			context['error'] = kwargs['error']
			context['message'] = kwargs['message']
		else:
			context['error'] = 'false'
			context['success'] = 'false'
			context['message'] = ' '

		context['supply_id'] = hardware.supply.id
		context['hardware_id'] = hardware.id 

		today = datetime.datetime.now()
		one_month = dateutil.relativedelta.relativedelta(months=1)
		a_month_ago = today - one_month

		context['today'] = today.strftime('%Y-%m-%d')
		context['a_month_ago'] = a_month_ago.strftime('%Y-%m-%d')
		# context['department'] = holder['department']

		if kwargs.get('message') is not None:
			context['error'] = kwargs.get('error')
			context['success'] = kwargs.get('success')
			context['message'] = kwargs.get('message')

		context['host_ip'] = 'http://' + socket.gethostbyname(socket.gethostname())
		return render(request, self.template, context)

	def post(self, request, *args, **kwargs):
		try:#Incase there is a error validating data or accessing db redirect to same hardware page.
			details = dict( request.POST )
			details = unpack(details)
			if details['purpose'] == 'update_hardware':
				return self.update_hardware( details )
			elif details['purpose'] == 'replace_hardware':
				return self.replace_hardware( request, details )
			elif details['purpose'] == 'new_incident':
				return self.new_incident( details )
			elif details.get('purpose') == 'resolve_incident':
				return self.resolve_incident(details)
			elif details.get('purpose') == 'assign_hardware':
				hardware = Hardware.objects.get(id=details['hardware_id'])
				hardware.assigned_to = details['staff_name']
				hardware.department = details['department']
				hardware.save()
				return self.get(request, **{'id': details['hardware_id'], 'type': hardware.hardware_type, 'message':'Hardware successfully assigned to ' + details['staff_name'], 'success':'true', 'error':'false' })
		except Exception as e:
			# print('\n\n\n\n\n\n\n' , e, '\n\n Incase there is a error validating data or accessing db redirect to same hardware page.')
			self.error = 1
			return redirect('/asset/viewAsset/' + request.POST['id'] + '/' + request.POST['hardware_type'])

class Supplied_hardware(View):
	template = 'Supplied_hardware.html'
	# Var to hold current hardware being handled if any error occurs. 
	error_hardware_holder = None

	def gen_report(request, *args, **kwargs):
		try:
			document = Document()
			width, height = document.sections[0].page_height, document.sections[0].page_width
			document.sections[0].orientation = WD_ORIENT.LANDSCAPE
			document.sections[0].page_width = width
			document.sections[0].page_height = height
			# To align title in center
			from docx.enum.text import WD_ALIGN_PARAGRAPH
			document.add_heading('ROKEL COMPUTER ASSETS REPORT \n' + str( datetime.datetime.now().strftime('%d-%m-%Y') ), 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
			flag = False
			details =  dict( request.POST )

			del details['csrfmiddlewaretoken']
		
			# # Check if branch exists in db if not remove so no values generated in report.
			# for  branch in details['generate_for']:
			# 	# if not bool ( Branch.objects.filter(name=branch) ) and branch != 'All Branches':
			# 	# 	del test[i]
			# 	# if not bool( Supply.objects.filter(branch=branch)):
			# 	details['generate_for'].remove(branch)
			# 	print(details['generate_for'])
			# 	# del test[i]
			# 	# print('ran for ', branch )
			# 	# document.add_paragraph(branch + " does not have any hardware supplied.")
			# 	# document.add_paragraph()

			details['generate_for'] = [branch for branch in details['generate_for'] if bool( Supply.objects.filter(branch=branch))]

			if ( len(details['generate_for']) < 1 ):
				flag = True
				raise Exception('The selected branches do not have any hardware supplied.')

			# Get all supplies for each branch.
			for branch in details['generate_for']: #because, for some reason, JS sent 'generate_for[]' in ajax instead of 'generate_for'.
				# Check if branch exists in db if not remove so no values generated in report.
				# if not bool ( Branch.objects.filter(name=branch) ) and branch != 'All Branches':
				# 	details['generate_for'].remove(branch)
				# 	pass
				# # If all branches are not in db halt report gen process by raising error.
				# if len (details['generate_for']) < 1:
				# 	raise Exception('Could not generate report for those branches.')
				document.add_heading(branch + " Branch", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
		
				hardware = []
				i = 0
				for indexs, supply in enumerate( Supply.objects.filter(branch=branch) ):
					#document.add_heading("Supply " + str(indexs+1), level=2)
					# document.add_paragraph("Date supplied : " + str(supply.date_of_supply.strftime('%d-%m-%Y')) )
					# document.add_paragraph( "Recipient staff : " + str(supply.assigned_to) )
					for index, hardware_component in enumerate( Hardware.objects.filter(supply=supply) ):
						i=i+1
						if hardware_component.hardware_type == "SystemUnit":
							document.add_heading("Hardware " + str(i) + " - WorkStation", level=2)
						else:
							document.add_heading("Hardware " + str(i) + " - " + hardware_component.hardware_type, level=2)
						tabs = "\t\t\t\t"
						document.add_paragraph("Department - " + hardware_component.department + tabs + "Assigned to - "+ hardware_component.assigned_to)
						# document.add_paragraph("Date supplied : " + str(supply.date_of_supply.strftime('%d-%m-%Y')) )

						# if supply.date_of_confirmation is None:
						# 	document.add_paragraph("Date received : Not Yet Received")
						# else:
						# 	document.add_paragraph("Date received : " + str(supply.date_of_confirmation.strftime('%d-%m-%Y')) )

						# if hardware_component.is_replaced is True:
						# 	replacement = Replacement.objects.get(old_hardware=hardware_component)
						# 	document.add_paragraph("This hardware component has been replaced by a " + replacement.new_hardware.hardware_type )	
						# Fields common to all hardware
						# if hardware_component.hardware_type == "SystemUnit":
						# 	document.add_paragraph("Hardware type - Work Station\t\t" + "Brand and model - " + hardware_component.brand_and_model)
						# else:
						# 	document.add_paragraph("Hardware type - " + hardware_component.hardware_type+"\t\tBrand and model - " +hardware_component.brand_and_model  )
						document.add_paragraph("Condition - " + hardware_component.condition + "\t\t\t\t\tBrand and model - " + hardware_component.brand_and_model )
						if hardware_component.hardware_type != 'SystemUnit':
							document.add_paragraph("Serial number - " + hardware_component.serial_number )
						elif hardware_component.hardware_type == 'SystemUnit' or hardware_component.hardware_type == 'Laptop':
							hardware_component = System_unit.objects.get(id=hardware_component.id)
							if hardware_component.ms_office_up_to_date is True:
								document.add_paragraph("PC name - "+ hardware_component.system_unit_name + "\t\t\tMS Office - Up to date")
							else:
								document.add_paragraph("PC name - "+ hardware_component.system_unit_name + "\t\t\tMS Office - Not up to date")
							
							document.add_paragraph("CPU serial number - " + hardware_component.serial_number)
							document.add_paragraph("Service tag code - " + hardware_component.service_tag_code )
							document.add_paragraph("Monitor serial number - " + hardware_component.monitor_serial_number)
							document.add_paragraph("Keyboard serial number - " + hardware_component.keyboard_serial_number )
							document.add_paragraph("Mouse serial number - " + hardware_component.mouse_serial_number)
							if hardware_component.printer_brand_and_model != '':
								document.add_paragraph('Assigned printer - ' + hardware_component.printer_brand_and_model + "\tPrinter serial number - "+hardware_component.printer_serial_number)
								# document.add_paragraph('Printer serial number : ' + hardware_component.printer_serial_number)
							if hardware_component.scanner_brand_and_model != '':
								document.add_paragraph('Assigned scanner - ' + hardware_component.scanner_brand_and_model + "Scanner serial number - " + hardware_component.scanner_serial_number)
								# document.add_paragraph('Scanner serial number : ' +  hardware_component.scanner_serial_number)
							if hardware_component.modem_brand_and_model != '':
								document.add_paragraph('Assigned modem - ' + hardware_component.modem_brand_and_model + "\tModem serial number - " + hardware_component.modem_serial_number)
								# document.add_paragraph('Modem serial number : ' + hardware_component.modem_serial_number)
							if hardware_component.antivirus_up_to_date is True:
								document.add_paragraph('Antivirus - Up to date\tOperating system - ' + hardware_component.operating_system)
							else:
								document.add_paragraph('Antivirus - Not up to date\tOperating system - ' + hardware_component.operating_system)

							if hardware_component.operating_system_up_to_date is True:
								document.add_paragraph('Operating system - Up to date\t')
							else:
								document.add_paragraph('Operating system - Not up to date\t')
						# Fields specific to software capable hardware.
						elif hardware_component.hardware_type == 'Server':
							hardware_component = Software_capable_hardware.objects.get(id=hardware_component.id)
							if hardware_component.antivirus_up_to_date is True:
								document.add_paragraph('Antivirus - Up to date')
							else:
								document.add_paragraph('Antivirus - Not up to date')

							document.add_paragraph('Operating system - ' + hardware_component.operating_system)

							if hardware_component.operating_system_up_to_date is True:
								document.add_paragraph('Operating system - Up to date')
							else:
								document.add_paragraph('Operating system - Not up to date')

							document.add_paragraph("Serial number - " + hardware_component.serial_number )
			f = io.BytesIO()
			document.save(f)
			f.seek(0)
			return FileResponse(f, as_attachment=True, filename='ROKEL COMPUTER ASSETS REPORT-' + datetime.datetime.now().strftime('%d-%m-%Y') +'.docx')
		except Exception as e:
			print('\n\n\n\n\n', e, '\n\n\n\n\n')
			if flag :
				return Supplied_hardware().get(request, **{'message': str(e), 'error':'true', 'success':'false'})
			else:
				return redirect('/asset/suppliedHardware')
	def get(self, request, *args, **kwargs):
		context = {'date':date, 'error':0}
		context['branch_name'] = request.user.username
		context['supply_hardware_form'] = Supply_hardware_form()
		context['supply_hardware_form'].fields.pop('date_of_supply')
		context['supply_hardware_form'].fields['department'] = forms.CharField(   widget=forms.Select( attrs=gen_attrs( ) ), label='Department or location')
		context['supply_hardware_form'].fields.pop('date_of_confirmation')
		context['supply_hardware_form'].fields.pop('branch')
		#Get all branches from django.auth
		from django.contrib.auth import get_user_model
		branches = get_user_model().objects.filter(  Q(is_superuser=0) & Q(is_staff=0) & ~Q(username='Admin') & ~Q(username='FHO') )
		#Gen dict of departments by branches e.g. { 'FHO':[ ['Risk','Risk'], ['Legal', 'Legal'] ]  , 'FCHAR':[  ['Finance','Finance']... ] }
		import json
		# context['all_departments_choices'] = {branch.username:[(department.name, department.name)for department in Department.objects.filter(branch=branch.username)] for branch in branches}
		# context['all_departments_choices'] =  json.dumps(context['all_departments_choices'])

		context['add_hardware_form'] = Add_hardware_form()

		context['Software_capable_hardware_form_components'] = Software_capable_hardware_form_components(initial={'operating_system':'Windows Server 2012 R2'})
		context['System_unit_form_components'] = System_unit_form_components(initial={'operating_system':'Windows 10 Pro'})
		context['System_unit_form_components'].fields['system_unit_name'].label = 'PC name'
		context['Other'] = Other() #Input field
		context['report_form'] = Report_form()
		branches = Branch.objects.filter(is_active=True)
		#Gen branch choices
		branches = gen_choices(query_set=branches, attr='name', field_name='Branch') #Choices is defined in forms.py
		holder = ('All Branches','All Branches')
		branches = [branch for branch in branches]
		branches.insert(1, holder)
		context['report_form'].fields['generate_for'].choices = branches
		context['report_form'].fields['generate_for'].widget.attrs['class'] = 'ui fluid dropdown'
		context['report_form'].fields['generate_for'].widget.attrs['multiple'] = ""

		if request.user.username != 'IT SUPPORT' and request.user.username != 'Admin':
			context['supplies'] = Supply.objects.filter( ~Q(status='System_Init'), branch=context['branch_name']) #system init is the status for hardware supplied before this system was established. 
		else:
			supplies = Supply.objects.filter( ~Q(status='System_Init') )
			context['supplies'] = supplies
		if bool(kwargs):
			context['message'] = kwargs['message']
			context['error'] = kwargs['error']
			context['success'] = kwargs['success']
		else:
			context['error'] = 'false';
			context['success'] = 'false';
			context['message'] = ' ';

		context['branches_to_select_from']= Branch.objects.all()

		return render(request, self.template, context)

	def post(self, request, *args, **kwargs):
		try:
			context = {'date':date, 'error':0}
			details = dict( request.POST )
			details['branch'] = details['branch'][0]
			details['assigned_to'] = details['assigned_to'][0]

			#See if hardware with serial number already exists. Prevents duplicate submissions and mistakes where s/n is the same.
			serial_numbers = []
			serial_numbers_str = ''
			for key, val in details.items():
				if 'serial_number' in key:
					serial_numbers.append(details[key][0])
			for val in serial_numbers:
				if bool ( Hardware.objects.filter(serial_number=val) ):
					return self.get(request,**{'error':'true','success':'false','message':'A Hardware Component With Serial Number: ' + val + ' Already Exists'})

			# Make supply if it does not exist.
			#[0] are need to sccess details since unpack is not run on the incoming request.POST obj.
			details['date_of_supply'] = datetime.datetime.now().strftime('%Y-%m-%d')

			try:#Check if supply exists,Make new one if not.
				supply = Supply.objects.get(branch=details['branch'], date_of_supply=details['date_of_supply'], assigned_to=details['assigned_to']) #Check if a supply with the same branch and date_of_supply already exists. If not create a new supply.

			except Exception as e:
				supply = Supply(branch=details['branch'], date_of_supply=details['date_of_supply'], assigned_to=details['assigned_to']) #Create new supply that hardware belongs to.
				supply.save()
		

			# Remove fields that belong to supply.
			details.pop('branch')
			details.pop('assigned_to')
			details.pop('date_of_supply')

			hardware = [] 

			# Make hardware components
			if int (details['num_SystemUnit'][0]) != 0:
				holder = {}
				i=1
				while i <= int (details['num_SystemUnit'][0]): 
					for key, val in details.items():
						if 'SystemUnit' + str(i) in key:
							holder[key.replace('SystemUnit'+str(i),'')] = details[key][0]
					i += 1
					holder['hardware_type'] = 'SystemUnit'
					hardware.append( System_unit(**holder) )
			if int (details['num_laptop'][0]) != 0:
				holder = {}
				i=1
				while i <= int (details['num_laptop'][0]): 
					for key, val in details.items():
						if 'laptop' + str(i) in key:
							holder[key.replace('laptop'+str(i),'')] = details[key][0]
					i += 1
					holder['hardware_type'] = 'Laptop'
					hardware.append( System_unit(**holder) )
			if int (details['num_server'][0]) != 0:
				holder = {}
				i=1
				while i <= int (details['num_server'][0]): 
					for key, val in details.items():
						if 'server' + str(i) in key:
							holder[key.replace('server'+str(i),'')] = details[key][0]
					i += 1
					holder['hardware_type'] = 'Server'

					holder.pop('Server')
					hardware.append( Software_capable_hardware(**holder) )

			if int (details['num_Other'][0]) != 0:
				holder = {}
				i=1
				while i <= int (details['num_Other'][0]): 
					for key, val in details.items():
						if 'Other' + str(i) in key:
							holder[key.replace('Other'+str(i),'')] = details[key][0]
						if '_type' in key:
							if details['_typeOther'+str(i)][0] == 'Laptop' :
								holder['hardware_type'] = 'Laptop    '
							else:
								holder['hardware_type'] = details['_typeOther'+str(i)][0]
					i += 1
					holder.pop('_type')
					hardware.append( Hardware(**holder) )

			for obj in hardware:#assign all hardware to supply and save.
				obj.supply = supply 
				obj.date_of_supply = supply.date_of_supply
				self.error_hardware_holder = obj
				obj.hardware = details['branch']
				obj.save()

			return redirect('/asset/suppliedHardware')
		except Exception as e:
			print('\n\n\n\n\n\n\n' , e)
			if str(e) == 'UNIQUE constraint failed: aMgnt_hardware.serial_number':
				return self.get(request, **{'message': self.error_hardware_holder.hardware_type + ' with serial number : ' + self.error_hardware_holder.serial_number + ' already exists.', 'error':'true'})
			return redirect('/asset/suppliedHardware')

class View_all_assets(View):
	template = 'view_all_assets.html'
	error = 0
	message = ''
	coming_from_form_submit = True
	def aMgntViewAllAssetsTable(request, *args, **kwargs):
		result = []
		query = request.POST['search[value]']
		query = query.split("|") 
		for index,word in enumerate(query):
		    query[index] = word.strip()
		start = int(request.POST['start'])
		length = int(request.POST['length'])
		query_set_holder = {}
		if request.user.username == 'IT SUPPORT' or request.user.username == 'Admin':
			if len(query) is 1:
				query = query[0]
				query = Q(branch__icontains=request.POST['search[value]']) | Q(hardware_type__icontains=query) | Q(brand_and_model__icontains=query) | Q(serial_number__icontains=query) | Q(condition__icontains=query) | Q(department__icontains=query) | Q(assigned_to__icontains=query) | Q(date_of_supply__icontains=query)
				query_set_holder = Hardware.objects.filter(query)
				hardware = query_set_holder[start:start+length]
			else:
				keywords = query
				for index, word in enumerate(keywords):
					if index == 0:
						query = Q(branch__icontains=word) | Q(hardware_type__icontains=word) | Q(brand_and_model__icontains=word) | Q(serial_number__icontains=word) | Q(condition__icontains=word) | Q(department__icontains=word) | Q(assigned_to__icontains=word) | Q(date_of_supply__icontains=word)
					else:
						query &= Q(branch__icontains=word) | Q(hardware_type__icontains=word) | Q(brand_and_model__icontains=word) | Q(serial_number__icontains=word) | Q(condition__icontains=word) | Q(department__icontains=word) | Q(assigned_to__icontains=word) | Q(date_of_supply__icontains=word)
				query_set_holder = Hardware.objects.filter(query)
				hardware = query_set_holder[start:start+length]
		else:
			if len(query) is 1:
				query = query[0]
				query = Q(hardware_type__icontains=query) | Q(brand_and_model__icontains=query) | Q(serial_number__icontains=query) | Q(condition__icontains=query) | Q(department__icontains=query) | Q(assigned_to__icontains=query) | Q(date_of_supply__icontains=query)
				query_set_holder = Hardware.objects.filter(query,branch=request.user.username)
				hardware = query_set_holder[start:start+length]
			else:
				keywords = query
				for index, word in enumerate(keywords):
					if index == 0:
						query = Q(branch__icontains=word) | Q(hardware_type__icontains=word) | Q(brand_and_model__icontains=word) | Q(serial_number__icontains=word) | Q(condition__icontains=word) | Q(department__icontains=word) | Q(assigned_to__icontains=word) | Q(date_of_supply__icontains=word)
					else:
						query &= Q(branch__icontains=word) | Q(hardware_type__icontains=word) | Q(brand_and_model__icontains=word) | Q(serial_number__icontains=word) | Q(condition__icontains=word) | Q(department__icontains=word) | Q(assigned_to__icontains=word) | Q(date_of_supply__icontains=word)
				query_set_holder = Hardware.objects.filter(query)
				hardware = query_set_holder[start:start+length]
		for hardware in hardware:
			if hardware.condition == 'Good':
				bs_class = 'success' 
			elif hardware.condition == 'Mid':
				bs_class = 'warning' 
			elif hardware.condition == 'Bad':
				bs_class = 'danger'
			if hardware.hardware_type == 'SystemUnit':
				hardware.hardware_type = 'WorkStation' 
			holder = {'DT_RowId':f'{hardware.id}/{hardware.hardware_type}/','Date of Supply':hardware.date_of_supply,'Department':hardware.department,'Assigned Role':hardware.assigned_to,'Serial Number':hardware.serial_number,'Brand and Model':hardware.brand_and_model,'Type':hardware.hardware_type,
					'Condition':f"<span class='bg-{bs_class} align-self-center' style='height:10px; width:10px; border-radius: 50%; display:inline-block;'></span>"}
			if request.user.username == 'IT SUPPORT' or request.user.username == 'Admin':
				holder['Branch'] = hardware.branch
			result.append(holder)
		return JsonResponse({'draw':int(request.POST['draw']),'recordsTotal':Hardware.objects.all().count(), 'recordsFiltered':query_set_holder.count(),'data':result})
	def get(self, request, *args, **kwargs):
		
		context = {'date':date, 'branch_name':request.user.username}

		if request.user.username != 'IT SUPPORT' and request.user.username != 'Admin': # Limit to hardware to that which belong to a branch.
			context['assets']  = [ Hardware.objects.filter(supply=supply, is_replaced=False, is_received=True) for supply in Supply.objects.filter(branch=context['branch_name']) ]
		else:# Get all hardware for IT SUPPORT and admin
			context['assets'] = []
			count = 0
			# for supply in Supply.objects.all():
			# 	for hardware in Hardware.objects.filter(supply=supply, is_replaced=False, is_received=True):
			# 		if count is not 100:
			# 			count += 1
			# 			context['assets'].append(hardware)
			# 		else:
			# 			break;
			# 	if count is 100:
			# 		break;
		# Gen message for user depending on previous input from post section of this view.
		context['success'] = 0
		context['error'] = 0 
		if bool( kwargs ):
			
			if kwargs['error'] == 'True':
				context['message'] = kwargs['message']
				context['show'] = 'true'
				context['error'] = 'true'
				context['success'] = 'false'
			elif kwargs['success'] == 'True':
				context['message'] = kwargs['message']
				context['show'] = 'true'
				context['success'] = 'true'
				context['error'] = 'false'
			
		else:
			context['show'] = 'false'
			context['message'] = ''

		if request.user.username == 'Admin':
			context['report_form'] = Report_form()
			branches = Branch.objects.filter(is_active=True)
			#Gen branch choices
			branches = gen_choices(query_set=branches, attr='name', field_name='Branch') #Choices is defined in forms.py
			context['report_form'].fields['generate_for'].choices = branches 

		###### Temporary function to add hardware for easy integration of system ########
		# Gen supply details form.
		context['supply_details_form'] = Supply_hardware_form()
		#context['supply_details_form'].fields['date_of_supply'].widget.attrs['value'] = datetime.datetime.now().strftime('%Y-%m-%d')
		context['supply_details_form'].fields['date_of_supply'].widget.attrs['max'] = datetime.datetime.now().strftime('%Y-%m-%d')
		context['supply_details_form'].fields['date_of_supply'].widget.attrs['min'] = '2000-01-01'
		#context['supply_details_form'].fields.pop('assigned_to')
		context['supply_details_form'].fields['department'] = forms.CharField(label='Location', widget=forms.Select(attrs=gen_attrs( {'class' :'form-control', 'required':'', 'placeholder':'e.g. Finance Department'}) ) )
		context['supply_details_form'].fields['department'].widget.choices = [(branch.name,branch.name) for branch in Branch.objects.filter(~Q(name="IT SUPPORT"),~Q(name="Admin"))] + [('RISK/ADVANCES', 'RISK/ADVANCES'),('ADMIN','ADMIN'),('OPERATIONS','OPERATIONS'),('COMPLIANCE','COMPLIANCE'), ('SAVINGS', 'SAVINGS'), ('CURRENT','CURRENT'), ('ENTRIES','ENTRIES'), ('TREASURY','TREASURY'), ('FINANCE','FINANCE'), ('COMPUTER CENTER','COMPUTER CENTER'), ('FOREIGN','FOREIGN'), ('MERCURY,', 'MERCURY')]
                # Gen department choices based on branch.
                #context['supply_details_form'].fields['department']
		# departments = Department.objects.filter(branch=context['branch_name'])
		# departments = [(department.name, department.name)for department in departments]
		# context['department_input'] = Department_form()
		# if not bool( departments ):
		# 	departments.append(('','No Departments or locations Have Been Added To This Branch'))
		# 	context['department_input'].fields['department'].widget.attrs['disabled'] = ''
		# 	context['deps'] = False
		# else:
		# 	context['deps'] = True
		# 	departments.insert(0, ('', '--Choose Department--') )
		# context['department_input'].fields['department'].widget.choices = departments
		context['supply_details_form'].fields.pop('branch')
		context['supply_details_form'].fields.pop('date_of_confirmation')
		context['supply_details_form'].fields['assigned_to'].label = 'Assigned To'
		context['supply_details_form'].fields.pop('assigned_to')
		context['supply_details_form'].fields['date_of_supply'].label = 'Date of Receipt'
		####################################################################################

		context['add_hardware_form'] = Add_hardware_form()
		# context['add_hardware_form'].fields['serial_number'].label = 'CPU / Serial number'
		context['add_hardware_form'].fields.pop('serial_number')
		context['Software_capable_hardware_form_components'] = Software_capable_hardware_form_components()
		context['System_unit_form_components'] = System_unit_form_components(initial={'operating_system':'Windows 10 Pro'})
		# context['System_unit_form_components'].fields['service_tag_code'].widget.attrs.pop('required')
		context['Other'] = Other() #Input field

		return render(request, self.template, context)
	def post(self, request, *args, **kwargs):
		
		# if 1==1:
		# 	return redirect('/asset/viewAllAssets/')
		try:
                        try:
                            for key, val in request.POST.items():
                                if 'serial_number' == key[:13]:
                                    if val == "NA":
                                            break
                                    elif bool ( Hardware.objects.filter(serial_number=val) ):
                                            raise Exception('Serial number : ' + str(val) + ' already exists.')
                                            break
                        except Exception as e:
                            return redirect('view_all_assets_success', message=str(e), success='False', error='True')
                        try:
                                supply = Supply.objects.get(assigned_to='System Init', date_of_supply=datetime.datetime.now().strftime('%Y-%m-%d'), branch=request.user.username)
                        except:
                                supply = Supply(assigned_to='System Init', date_of_supply=datetime.datetime.now().strftime('%Y-%m-%d'), date_of_confirmation=datetime.datetime.now().strftime('%Y-%m-%d'), branch=request.user.username)
                                # supply.status = 'Received'
                                supply.status = 'System_Init'
                                supply.save()

                        for num in range(int (request.POST['num_hardware']) ):
                              
                                hardware = dict()
                                hardware['supply'] = supply
                                hardware['is_received'] = True
                                #hardware['assigned_to'] = request.POST['assigned_to']
                                #hardware['department'] = request.POST['department']
                                for key, val in request.POST.items():
                                        if str(num+1) in key:
                                                hardware[key.replace(str(num+1),'')] = val

      
                                if hardware['hardware_type'] == "Other":
                                	hardware['hardware_type'] = request.POST['_type1']
                                	hardware.pop('_type')
                                	holder = Hardware(**hardware)
                                	holder.branch = request.user.username
                                	holder.save()
                                elif hardware['hardware_type'] == "Server":
                                    holder = Software_capable_hardware(**hardware)
                                    holder.branch = request.user.username
                                    holder.save() 
                                elif hardware['hardware_type'] == "SystemUnit":
                                    holder = System_unit(**hardware)
                                    holder.branch = request.user.username
                                    holder.save() 
                                elif hardware['hardware_type'] == "Laptop":
	                                holder = System_unit(**hardware)
	                                holder.branch = request.user.username
	                                holder.save() 
                            #     elif hardware['hardware_type'] == "Laptop":
		                        		# System_unit(**hardware).save()      
                        return redirect('view_all_assets_success', message = 'Hardware component(s) added successfully', error = 'False', success = 'True')
		except Exception as e:
			print('\n\n\n', e, '\n\n\n\n')
			return redirect('/asset/viewAllAssets/')


		# print(request.POST)
		# if True:
		# 	return HttpResponse('.workgin.')
		# try:
		# 	purpose = request.POST.get('purpose')
		# 	details =  dict(request.POST)
		# 	unpack(details)

		

		# 	elif purpose == 'add_hardware': #This section is only for easy integration. Remove when all existing hardware has been inputted.
		# 		details = unpack(request.POST)
		# 		details.pop('purpose')
		# 		if details.get('hardware_type') is not None: #User hasn't tampered with form.
		# 			try:
		# 				supply = Supply.objects.get(branch=request.user.username, date_of_supply=details['date_of_supply'])
		# 			except:
		# 				supply = Supply(assigned_to=details['assigned_to'],
		# 								branch = request.user.username,
		# 								date_of_confirmation = datetime.datetime.now().strftime('%Y-%m-%d'), 
		# 								date_of_supply = details['date_of_supply'], 
		# 								status='System_Init'
		# 							)

		# 				supply.save()

		# 			details.pop('date_of_supply')
		# 			details.pop('assigned_to')

		# 			try:#See if hardware with serial number already exists. Prevents duplicate submissions and mistakes where s/n is the same.
		# 				Hardware.objects.get(serial_number=details['serial_number'])
		# 				return self.get(request,kwargs={'error': 1, 'message':'A Hardware Component With That S/N Already Exists'})
		# 			except:
		# 				if System_unit_form_components(details).is_valid():
		# 					hardware = System_unit(**details)
		# 				elif Software_capable_hardware_form_components(details).is_valid():
		# 					hardware = Software_capable_hardware(**details)
		# 				elif Other(details).is_valid():
		# 					# Get hardware type from _type and remove _type.
		# 					details['hardware_type'] = details['_type']
		# 					details.pop('_type') 
		# 					hardware = Hardware(**details)

		# 				hardware.supply = supply # All inital hardware get assigned to supply id 0.
		# 				hardware.is_received = True
		# 				hardware.is_replaced = False
		# 				hardware.save()

		# 		return redirect('/asset/viewAllAssets')

		# 	elif purpose == 'register_department':

				
		# 		if not bool ( Department.objects.filter( name=request.POST.get('department'), branch=request.user.username ) ): #check if department already exists.
		# 			Department(name=request.POST['department'], branch=request.user.username).save()

		# 			return self.get(request, kwargs={'success':True, 'message':'New Department Succesfully Added'})

		# 		else:

		# 			return self.get(request, kwargs={'error':True, 'message':'Department Already Exists'})

				
		# 	else:
		# 		return redirect('/asset/viewAllAssets/')
		# except Exception as e:
		# 	print(e, '\n\n\n\nfrom here')
		# 	return redirect('/asset/viewAllAssets/')
class Supply_details(View):
	template = 'supply_details.html'

	def get(self, request, *args, **kwargs):
		context = {'date':date, 'branch_name':request.user.username}

		supply = Supply.objects.get(id=kwargs['id'])
		supplied_hardware = Hardware.objects.filter(supply=supply)
		list_of_hardware = []

		if kwargs.get('error') is not None:
			context['message'] = kwargs['message']
			if kwargs.get('error') is not None:
				context['error'] = kwargs['error']
			else:
				context['error'] = 'true'
			if kwargs.get('success') is not None:
				context['success'] = kwargs['success'] 
			else:
				context['success'] = 'false'
		elif kwargs.get('success') is not None:
			context['error'] = 'false'
			context['success'] = 'true'
			context['message'] = kwargs['message']
		else:
			context['error']= 'false'
			context['success']= 'false'
			context['message'] =  '' 

		#supply = {

		#{   'department'  : 'Finance Department' ,  'assigned_to': 'Finance Manager' ,  list_of_supplied_hardware : [Hardware_obj(1),Hardware_obj(209)...]   },

		#{   'department'  : 'Risk Department' ,  'assigned_to': 'Risk Manager' ,  list_of_supplied_hardware : [Hardware_obj(31),Hardware_obj(24)...]   },

		#{   'department'  : 'Marketing Department' ,  'assigned_to': 'Marketing Manager' ,  list_of_supplied_hardware : [Hardware_obj(12),Hardware_obj(25)...]   },

		#...         

		#}

		for hardware in supplied_hardware:
			if hardware.hardware_type == 'Server':
				list_of_hardware.append ( Software_capable_hardware.objects.get(id=hardware.id).__dict__ )
			elif hardware.hardware_type == 'SystemUnit' or hardware.hardware_type == 'Laptop':
				list_of_hardware.append ( System_unit.objects.get(id=hardware.id).__dict__ )
			else:
				list_of_hardware.append ( hardware.__dict__ )

		context['list_of_hardware'] = list_of_hardware
		context['sup_id'] = supply.id
		context['supply'] = supply
		context['host_ip'] = 'http://' + socket.gethostbyname(socket.gethostname())
		return render(request, self.template, context)

	def post(self, request, *args, **kwargs):
		try:
			details = unpack(request.POST)
			#### For now a sing receipt button is used to confirm entire supply.#####
			########## Section for validating each hardware part of supply in a checklist fashion. Implement later#############
			# hardware = Hardware.objects.get(id=details['id'])
			# hardware.is_received = True
			# hardware.save()
			# Check if any hardware part of the supply is not yet ackowledge. This is for if we add a confirm option for each hardware in supply.
			# is_confirmed = Hardware.objects.filter( is_received=False, supply=Supply.objects.get( id=int (details['sup_id']) ) )

			# if not bool (is_confirmed):#check if there are any non-confirmed hardware part of supply.
			# 	supply = Supply.objects.get( id=int(details['sup_id']) )
			# 	supply.status = 'Confirmed'#Confirm supply
			# 	supply.date_of_confirmation = datetime.datetime.now().strftime('%Y-%m-%d')
			# 	supply.save()
			###############################################################################################################
			if details['purpose'] == 'confirm_hardware':
				user = authenticate(username=request.user.username, password=details['branch_password'])#check if branch password is correct
				if user is not None:
					supply = Supply.objects.get( id=int(details['sup_id']) )
					supply.status = 'Received'#Confirm supply
					supply.date_of_confirmation = datetime.datetime.now().strftime('%Y-%m-%d')
					supply.save()
					# Set all hardware supplied is_received to True
					for hardware in Hardware.objects.filter(supply=supply):
						hardware.is_received = True
						hardware.save()

					return self.get(request, **{'id':details['sup_id'], 'success':'true', 'message':'Confirmed Successfully.'} )

				else:
					return self.get(request, **{'id':details['sup_id'], 'error':'true', 'message':'Password is incorrect.'} )

			elif details['purpose'] == 'assign_hardware':
				hardware = Hardware.objects.get(id=details['hardware_id'])
				hardware.assigned_to = details['staff_name']
				hardware.department = details['department']
				hardware.save()
				return self.get(request, **{'id': details['sup_id'], 'message':'Hardware successfully assigned to ' + details['staff_name'], 'success':'true', 'error':'false' })
			elif details['purpose'] == 'reject_supply':
				supply = Supply.objects.get(id=details['sup_id'])
				supply.status = 'Problem'
				supply.save()
				return self.get(request, **{'id': details['sup_id'], 'success': 'true', 'error': 'false',  'message':'This supply has been sent to IT SUPPORT for revision. This page will be updated after they make changes.'})
			elif details['purpose'] == 'edit_supply':
				hardware = Hardware.objects.get(id=details['hardware_id'])
				hardware.brand_and_model = details['Brand and model']
				hardware.serial_number = details['Serial number']
				hardware.save()
				return self.get(request, **{'id': details['sup_id'], 'success': 'true', 'error': 'false',  'message':'The selected hardware has been updated.'})
			elif details['purpose'] == 'resend_supply':
				supply = Supply.objects.get(id=details['supply_id'])
				supply.status = 'Pending' 
				supply.save()
				return self.get(request, **{'id': details['supply_id'], 'success': 'true', 'error': 'false',  'message':'This supply has been resent to branch ' + supply.branch})
			elif details['purpose'] == 'delete_supply':
				if authenticate(username=request.user.username, password=request.POST['admin_password']) is None:
					return self.get(request, **{'id': details['sup_id'], 'success': 'false', 'error': 'true',  'message':'Incorrect admin password.'})
				supply = Supply.objects.get(id=int(details['sup_id']))
				supply.delete()
				return Supplied_hardware().get(request, **{'success':'true', 'error':'false', 'message':'Supply for ' + supply.branch + ' on ' + supply.date_of_supply.strftime('%B %e, %Y') + ' successfully deleted.'})
			return redirect('/asset/suppliedHardware/')
		except Exception as e:
			# print("\n\n\n\n", e, '\n\n\n\n\n')
			return redirect('/asset/suppliedHardware/')

	def print_supply(request, *args, **kwargs):
		# try:
		details = dict (request.POST)
		details = unpack(details)
		if details['purpose'] == 'print_supply':
			supply = Supply.objects.get(id=details['sup_id'])
			hardwares = Hardware.objects.filter(supply=supply)
			hardware_types = {'server' : 0, 'desktop' : 0}
			for hardware in hardwares:
				if hardware.hardware_type == 'Server':
					hardware = Software_capable_hardware.objects.get(id=hardware.id)
					hardware_types.update({'server': hardware_types['server'] + 1})
				elif hardware.hardware_type == 'SystemUnit':
					hardware = System_unit.objects.get(id=hardware.id)
					hardware_types.update({'desktop': hardware_types['desktop'] + 1})
				else:
					if hardware_types.get(hardware.hardware_type.lower()) is None:
						hardware_types.update({hardware.hardware_type.lower() : 1})
					else:
						hardware_types.update({hardware.hardware_type.lower(): hardware_types[hardware.hardware_type.lower()] + 1})


			is_plural = 0
			for key, val in hardware_types.items():
				is_plural = is_plural + val
			# if more than one hardware in supply check if all are of the same type.
			all_the_same = False
			for key, val in hardware_types.items():
				if val > 0:
					all_the_same += 1
				if all_the_same > 1:
					all_the_same = False
					break

			
			all_the_same = True

			p = inflect.engine()
			document = Document()

			document.add_picture('./aMgnt/static/aMgnt/assets/logo/RokelLogoHeader.jpg', width=Inches(1.0), height=Inches(1.0))

			document.add_paragraph(supply.assigned_to)
			document.add_paragraph(str (supply.branch + " Branch") )
			document.add_paragraph("[Insert Branch Location Here]")
			document.add_paragraph("Your Ref: " + supply.branch + "\t\t\t\t\t\t Our Ref: AL/[Insert Staff Initials Here]")
			document.add_paragraph()
			document.add_paragraph( 'Date: ' + str( supply.date_of_supply.strftime('%d-%B-%Y') ).replace('-',' ') )
			document.add_paragraph( 'Subject: Computers and Accessories')
			document.add_paragraph()
			if is_plural > 1:
				to_be_conjugation = 'have' 
			else:
				to_be_conjugation = 'has'

			letter_body = ""
			index = 1
			for key, val in hardware_types.items():
			
				# make word plural e.g. server => servers, desktop => desktops, mouse => mice etc...
				if key == "mouse" and val > 1:
					key = "mice"
				elif val > 1:
					key = key + 's'
				if val != 0:
					letter_body += str(val) + " (" + p.number_to_words(val) + ") "+ key
				if is_plural > 1 and not all_the_same:
					if is_plural - index == 2:
						letter_body += ", "
					elif index != is_plural:
						letter_body += " and " 

				index += 1

			document.add_paragraph('We advise that ' + letter_body+ " " + to_be_conjugation +" been supplied to your branch." )
			document.add_paragraph()
			document.add_paragraph("Description \t \t \t \t \t \t  Serial No")
			for hardware in hardwares :
				if hardware.hardware_type != "SystemUnit":
					document.add_paragraph(hardware.hardware_type.upper() + "\t \t \t \t \t \t  CN=" + hardware.serial_number)
				else:
					hardware = System_unit.objects.get(serial_number=hardware.serial_number)
					document.add_paragraph("CPU\t \t \t \t \t \t \t  CN=" + hardware.serial_number)
					document.add_paragraph("MONITOR\t \t \t \t \t \t  CN=" + hardware.monitor_serial_number)
					document.add_paragraph("KEYBOARD\t \t \t \t \t \t  CN=" + hardware.keyboard_serial_number)
					document.add_paragraph("MOUSE\t \t \t \t \t \t  CN=" + hardware.mouse_serial_number)

			document.add_paragraph("\n\n\nPlease acknowledge receipt\n\n")
			document.add_paragraph("\n\nAlhaji  Lewally\nIT Manager")

			f = io.BytesIO()
			document.save(f)
			length = f.tell()
			f.seek(0)
			
			return FileResponse(f, as_attachment=True, filename="Supply For "+ supply.branch  +  ' ' + str(supply.date_of_supply) +".docx")
		return redirect('/asset/suppliedHardware/')
		# except Exception as e:
		# 	print('\n\n\n\n\n\n', e, '\n\n\n\n')
		# 	return redirect('/asset/suppliedHardware/')
class Create_supply(View):
	template = 'create_supply.html'

	def get(self, request, *args, **kwargs):
		context = {'date':date, 'branch_name':request.user.username}

		return render(request, self.template, context)

	def post(self, request, *args, **kwargs):

		return HttpResponse('under const..')

class aMgntGeneralIncidents(View):
	template = 'aMgntGeneralincidents.html'
	def get(self, request, *args, **kwargs):
		if bool(kwargs.get('message')): #check if there is a message to display.
			message = kwargs['message']
			success = kwargs['success']
		else:
			message = ''
			success = '-1'
		return render(request,self.template,{'date':date, 'generalIncidents':bool(GeneralIncident.objects.all()), 'branches':Branch.objects.filter(~Q(name="IT SUPPORT")), 'message':message,'success':success})
	def post(self, request, *args, **kwargs):
		try:
			if request.user.username == 'Admin' or request.user.username == 'IT SUPPORT':
				branch = request.POST['Branch']
				Branch.objects.get(name=branch)
			else:
				branch = request.user.username
			if request.POST.get('Resolution') is not None:
				resolution = request.POST['Resolution']
			else:
				resolution = ''
			new_incident = GeneralIncident(
								branch=branch,
								date=make_aware(datetime.datetime.strptime(request.POST['Date'] + ' ' + request.POST['Time'],'%Y-%m-%d %H:%M:%S')),
								description=request.POST['Description'],
								reportingStaffName=request.POST['ReportingStaffName'],
								severity=request.POST['Severity'],
								resolved=int(request.POST['Status']),
								resolution=resolution,
							)
			new_incident.save()
			return JsonResponse({'success':True,'message':'Incident reported successfully.','newIncidentId':new_incident.id})
		except Exception as e:
			log("\n\n\n" + str(e) + "\n\n\n")
			return JsonResponse({'success':False, 'message':'Could not add incident.'}) 
	def updateIncident(request):
		try:
			if request.method == 'POST':
				incident = GeneralIncident.objects.get(id=request.POST['id'])
				if incident.resolved is not 1: #not resolved.
					incident.resolution = request.POST['resolutionDetails']
					incident.resolved = 1
					incident.save()
				else:
					return JsonResponse({'success':False, 'message':'Incident has already been resolved.'})
				return JsonResponse({'success':True,'message':'Incident successfully updated.'})
			else:
				response = HttpResponse()
				response.status_code = 403
				return response
		except Exception as e:
			log('\n\n\n' + str(e) + '\n\n\n')
			return JsonResponse({'success':False,'message':'Error updating incident.'})
	def get_incidents(request):
		try:
			query = request.POST['search[value]']
			query = Q(branch__contains=query) | Q(reportingStaffName__icontains=query)
			if request.user.username == 'Admin' or request.user.username == 'IT SUPPORT':
				general_incidents = GeneralIncident.objects.filter(query).order_by('-date')[int(request.POST['start']):int(request.POST['start']) + int(request.POST['length'])]
			else:
				general_incidents = GeneralIncident.objects.filter(query,branch=request.user.username).order_by('-date')[int(request.POST['start']):int(request.POST['start']) + int(request.POST['length'])]
			if bool(general_incidents):
				if request.user.username == 'Admin' or request.user.username == 'IT SUPPORT':
					rows = [{'DT_RowId':incident.id,'Branch':incident.branch,'Date':incident.date.strftime('%A %d. %B %Y'),'Time':incident.date.strftime('%H:%m:%S'),'Description':incident.description,'Reporting Staff':incident.reportingStaffName,'Severity':incident.get_severity(),'Status':incident.get_status()}for incident in general_incidents]
				else:
					rows = [{'DT_RowId':incident.id,'Date':incident.date.strftime('%A %d. %B %Y'),'Time':incident.date.strftime('%H:%m:%S'),'Description':incident.description,'Reporting Staff':incident.reportingStaffName,'Severity':incident.get_severity(),'Status':incident.get_status()}for incident in general_incidents]
				return JsonResponse({'draw':int(request.POST['draw']),'recordsTotal':general_incidents.count(), 'recordsFiltered':GeneralIncident.objects.all().count(),'data':rows})
			else:
				return JsonResponse({'recordsTotal':0,'recordsFiltered':0,'data':[]})
		except Exception as e:
			log('\n\n\n' + str(e) + '\n\n\n')
			return JsonResponse({'success':False,'message':'Error retrieving incidents.','data':[]})
	def get_incident(request):
		try:
			pass
		except Exception as e:
			log('\n\n\n' + str(e) + '\n\n\n')
			response = HttpResponse()
			response.status_code = 500
			return response
class aMgntGeneralIncident(View):
	def get(self, request, *args, **kwargs):
		try:
			if bool(kwargs.get('message')): #check if there is a message to display.
				message = kwargs['message']
				success = kwargs['success']
			else:
				message = ''
				success = '-1'
			incident = GeneralIncident.objects.get(id=int(kwargs['incidentId']))
			return render(request,'aMgntViewIncident.html',{'date':date,'incident':incident,'success':success,'message':message})
		except Exception as e:
			log('\n\n\n'+str(e)+'\n\n\n')
			return aMgntGeneralIncidents().get(request, message="Could not load incident details", success='0')
	def post(self, request, *args, **kwargs):
		try:
			incident = GeneralIncident.objects.get(id=request.POST['id'])
			incident.resolution = request.POST['resolutionDetails']
			incident.resolved = 1
			incident.save()
			return render(request,'aMgntViewIncident.html',{'date':date,'incident':incident, 'success':'1', 'message':'Incident Resolved Successfully.'})
		except Exception as e:
			log('\n\n\n'+str(e)+'\n\n\n')
			return render(request,'aMgntViewIncident.html',{'date':date,'incident':incident, 'success':'0', 'message':'Incident Could Not Be Resolved.'})
	def delete_incident(request):
		try:
			user = authenticate(username='Admin', password=request.POST.get('adminPassword') )
			if user is None:
				return aMgntGeneralIncident().get(request,incidentId=request.POST['id'],success="0",message='Incorrect admin password.')
			incident = GeneralIncident.objects.get(id=request.POST['id'])
			incident.delete()
			return aMgntGeneralIncidents().get(request,success='1',message='Incident for '+incident.branch+' has been deleted.' )
		except Exception as e:
			log('\n\n\n' + str(e) + '\n\n\n')
			return aMgntGeneralIncident().get(request,success="0",message="Could not delete incident.")

class aMgntQRScanner(View):	
	def get(self,request,*args,**kwargs):
		return render(request,'aMgntQRScanner.html',{'date':date})
	def post(self,request,*args,**kwargs):
		try:
			hardware = Hardware.objects.get(id=request.POST['id'])
			return JsonResponse({'success':True,'hardwareDetails':{'supplyId':hardware.supply.id,'status':hardware.supply.status,'id':hardware.id,'Branch':hardware.supply.branch,'Department':hardware.department,'SerialNo':hardware.serial_number,'HardwareType':hardware.hardware_type}
})
		except Exception as e:
			log('\n\n\n'+str(e)+'\n\n\n')
			return JsonResponse({'success':False})

#for supply in Supply.objects.all():
#    for hardware in Hardware.objects.filter(supply=supply):
# 	    hardware.branch = supply.branch
# 	    hardware.save()

