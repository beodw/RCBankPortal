######################## Utility methods ######################
import datetime
date = datetime.datetime.now().strftime('%Y')
#################  Dependecies for core app  ####################
import json
from .forms import *
from .models import *
from django.views import View
from django.db.models import Q
from django.shortcuts import render
from django.shortcuts import redirect
from django.contrib.auth.models import User
from django.contrib.auth import authenticate, login
from django.http import FileResponse, JsonResponse, HttpResponse
from django.contrib.auth.hashers import make_password, check_password
############# Dependencies for genreating dashboard figures as excel file#########
import io
from docx import Document
from docx.table import Table
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from cherrypy import log
#############################
def nMeetLogin(request):
	try:
		if request.method == 'GET': #serve login page
			if request.user.is_authenticated : #except if already logged_in
				return redirect('dashboard')
			else :
				context = {'message':'', 'error':'none', 'date':date}
				return render(request, "nMeetLogin.html", context)
		if request.method == 'POST': #validate user
			user = authenticate(username=request.POST.get('username'), password=request.POST.get('password') )
			if user is not None :
				login(request, user, backend=user.backend)
				requested_page = request.GET.get('next')
				if requested_page is not None :
					return redirect(requested_page)
				else :
					return redirect('nMeetDashboard')
			else:
				context = {'message' : 'Incorrect Username or Password', 'error':'visible', 'date':date}
				return render(request, 'nMeetLogin.html',  context)
	except Exception as e: #catch any errors just to be extra safe.
		return redirect('nMeetLogin')
def nMeetLogout(request):
	from django.contrib.auth import logout
	logout(request)
	return redirect('nMeetLogin')
def http_response(status_code):
	response = HttpResponse()
	response.status_code = status_code
	return response
class nMeetViewDashboard(View):
	def get(self, request, *args, **kwargs):
		#Return the ui for the dashboard page.
		try:
			print(request.user.has_perm('nMeet.add_dashboard'))
			return render(request,'nMeetDashboard.html',{'date':date,'dashboard':True})
		#Incase error accessing db.
		except Exception as e:
			log('\n\n\n'+str(e)+'\n\n\n')
			return render(request,'nMeetDashboard.html',{'date':date,'dashboard':False})
	def post(self, request, *args, **kwargs):
		#Returns n newest dashboard entries when ajax request made. n = request.POST['length'].
		try:
			dashboard = []
			if request.POST['length'] == "-1":
				dashboard_buffer = Dashboard.objects.all()
			else:
				dashboard_buffer = Dashboard.objects.all()[:int(request.POST['length'])]
			for row in dashboard_buffer:
				if row.updated is True:
					time_sent = row.updated_at.strftime('%H:%M:%S')
					dashboard.append([row.id,row.branch,row.teller_or_operator,row.debit,row.credit,time_sent,"updated"])
				else:
					time_sent = row.created_at.strftime('%H:%M:%S')
					dashboard.append([row.id,row.branch,row.teller_or_operator,row.debit,row.credit,time_sent,"not_updated"]) 
			return JsonResponse({'dashboard':dashboard,'cancelled':[cancelled.id_of_cancelled_fig for cancelled in Cancelled.objects.all()]})
		except Exception as e:
			log('\n\n\n\n'+str(e)+'\n\n\n\n')
			return http_response(status_code=500)
def nMeetClearDashboard(request):
	try:
		if authenticate(username='Admin', password=request.POST['admin_password']) is None:
			return JsonResponse({'success':False})
		else:
			for row in Dashboard.objects.all():
				row.delete()
			for cancelled in Cancelled.objects.all():
				cancelled.delete()
			return JsonResponse({'success':True})
	except Exception as e:
		log('\n\n\n'+str(e)+'\n\n\n')
		return JsonResponse({'success':False})
def nMeetSaveDashboard(request):
	try:
		document = Document()
		from docx.enum.text import WD_ALIGN_PARAGRAPH
		document.add_heading('ROKEL END OF DAY PROCESSING \n' + str( datetime.datetime.now().strftime('%d-%m-%Y') ), 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
		query = ~Q(username="Admin") & ~Q(username="beod") & ~Q(username="IT SUPPORT") 
		users = User.objects.filter(query)
		branches_in_order = ['FSSST','FCLOCK','BO','KENEMA','FWIA','FHO','FWILF','KOIDU','MOYAMBA','PUJEHUN','FCX','MAKENI','FCHAR','KABALA','KAILAHUN','FWILK', 'FJUBA']
		for user in users:
			if user.username not in branches_in_order:
				branches_in_order.append(user.username)
		for branch_name in branches_in_order:
			branch = branch_name 
			document.add_heading(branch + ' Figures').alignment = WD_ALIGN_PARAGRAPH.CENTER
			rows = Dashboard.objects.filter(branch=branch)
			if bool(rows):
				table = document.add_table(len(rows)+1,4)
				table.cell(0,0).add_paragraph('Teller or Operator')
				table.cell(0,1).add_paragraph('Debit')
				table.cell(0,2).add_paragraph('Credit')
				table.cell(0,3).add_paragraph('Time Sent')
				table.alignment = WD_TABLE_ALIGNMENT.CENTER
				for index,row in enumerate(rows):
					cells = table.row_cells(index+1)
					for i,cell in enumerate(cells):
						if i == 0:
							cell.add_paragraph(row.teller_or_operator)
						elif i == 1:
							cell.add_paragraph(row.debit)
						elif i == 2:
							cell.add_paragraph(row.credit)
						elif i == 3:
							if row.updated_at != None:
								cell.add_paragraph(row.updated_at.strftime('%H:%M:%S'))
							else:
								cell.add_paragraph(row.created_at.strftime('%H:%M:%S')) 
			else:
				document.add_paragraph('No figures sent for ' + branch).alignment = WD_ALIGN_PARAGRAPH.CENTER
		f = io.BytesIO()
		document.save(f)
		f.seek(0)
		return FileResponse(f, as_attachment=True, filename='ROKEL END OF DAY PROCESSING FIGURES-' + datetime.datetime.now().strftime('%d-%m-%Y') +'.docx')
	except Exception as e:
		log('\n\n\n\n' + str(e) + '\n\n\n\n')
		return redirect('generic500page',redirect_url='error')
class nMeetEnterFigures(View):
	template = 'nMeetEnterFigures.html'
	def get(self, request, *args, **kwargs):
		return render(request,self.template,{'date':date})
	def post(self, request, *args, **kwargs):
		try:
			already_existing_row = Dashboard.objects.filter(branch=request.user.username,teller_or_operator=request.POST.get('teller_or_operator'))
			if bool ( already_existing_row ):
				already_existing_row[0].credit = request.POST.get('credit').replace(" ","").replace("\t","")
				already_existing_row[0].debit = request.POST.get('debit').replace(" ","").replace("\t","")
				already_existing_row[0].updated = True
				already_existing_row[0].save()
			else:
				row = Dashboard(branch=request.user.username,teller_or_operator=request.POST.get('teller_or_operator'),debit=request.POST.get('debit').replace(" ","").replace("\t",""),credit=request.POST.get('credit').replace(" ","").replace("\t",""))
				row.save()
			response = HttpResponse()
			response.status_code = 200
			return response
		except Exception as e:
			log('\n\n\n'+str(e)+'\n\n\n')
			response = HttpResponse()
			response.status_code = 500
			return response
class nMeetChat(View):
	def get(self, request, *args, **kwargs):
		chat = [[message.created_at.strftime('%Y-%m-%d %H:%M:%S'),message.branch,message.text,message.id] for message in Chat.objects.all()]
		return JsonResponse({'chat':chat})
	def post(self, request, *args, **kwargs):
		#Add chat and return success message.
		try:
			if request.POST.get('text') is not None and request.POST.get('text') != '':
				chat = Chat(branch=request.user.username, text=request.POST.get('text'))
				chat.save()
			return http_response(status_code=200)
		except Exception as e:
			log('\n\n\n'+str(e)+'\n\n\n')
			return http_response(status_code=503)
def refreshChat(request):
	#chat_buffer holds the last 50 rows of the chat.
	chat_buffer = [[message.created_at.strftime('%Y-%m-%d %H:%M:%S'),message.branch,message.text,message.id] for message in Chat.objects.all().order_by('id')[:50]]
	return JsonResponse({'chat_buffer':chat_buffer})
def nMeetClearChat(request):
	for message in Chat.objects.all():
		message.delete()
	for deleted_message in DeletedMessage.objects.all():
		deleted_message.delete()
	response = HttpResponse()
	response.status_code = 200
	return response
def get_deleted_messages(request):
	deleted_messages_list = [message.message_id for message in DeletedMessage.objects.all()]
	return JsonResponse({'deleted_messages':deleted_messages_list})
def deleteMessage(request):
	if request.method == 'POST':
		try:
			deleted_message = Chat.objects.get(id=int(request.POST['message_id']))
			deleted_message.delete()
			deleted_message = DeletedMessage(message_id=request.POST['message_id'])
			deleted_message.save()
			return JsonResponse({'success':1})
		except Exception as e:
			log('\n\n\n' + str(e) + '\n\n\n\n')
			return JsonResponse({'success':0})
	else:
		return JsonResponse({'success':0})
class nMeetCancelFigs(View):
	template = 'cancelFigs.html'
	branches = User.objects.filter(~Q(username="Admin") & ~Q(username="beod") & ~Q(username="IT SUPPORT"))
	def get(self,request,*args,**kwargs):
		try:
			if request.user.username == 'Admin' or request.user.username == 'IT SUPPORT':  
				return render(request,self.template,{'date':date, 'error':'false', 'success':'false', 'message':'', 'branches':self.branches})
			else:
				return redirect('nMeetDashboard')
		except Exception as e:
			log('\n\n\n'+ str(e)+ '\n\n\n')
			return render(request,self.template,{'date':date, 'error':'true', 'success':'false','message':'An error was encountered loading the page. Please contact admin.'})
	def post(self,request,*args,**kwargs):
		try:
			row = Dashboard.objects.filter(branch=request.POST['branch'],teller_or_operator=request.POST['teller_or_operator'])
			if bool(row):
				cancelled = Cancelled(id_of_cancelled_fig=row[0].id)
				cancelled.save()
				row[0].delete()
				return render(request, self.template, {'branches':self.branches,'date':date, 'success':'true', 'error':'false','message':'The figs for '+ str(row[0].branch)+ ' ' + str(row[0].teller_or_operator) +' have been deleted successfully.'})
			else:
				return render(request,self.template,{'branches':self.branches,'date':date,'success':'false','error':'true','message':'There are no figs for that teller / operator.'})
		except Exception as e:
			log('\n\n\n'+str(e)+'\n\n\n')
			return render(request, self.template, {"date":date, 'error':'true', 'message':'There was a problem deleting the selected figures. Please contact admin or IT CENTER for support.', 'success':'false'})
def nMeet_upload_summary_file(request):
	try:
		if request.method == "POST":
			figures = json.loads(request.POST['tellers_and_operators'])
			for teller_or_operator in figures: 
				payload = Dashboard.objects.filter(
						branch=request.user.username,
						teller_or_operator= teller_or_operator
				)
				if bool(payload):
					payload = payload.first() 
					payload.debit = figures[teller_or_operator]['debit']
					payload.credit = figures[teller_or_operator]['credit']
					payload.updated = True
					payload.save()
				else:
					payload = Dashboard(
						branch=request.user.username, 
						teller_or_operator=teller_or_operator,
						debit=figures[teller_or_operator]['debit'],
						credit=figures[teller_or_operator]['credit']
					)
					payload.save()
			return JsonResponse({'success':True, 'message':'Figures uploaded successfully.'})
		else:
			response = HttpResponse()
			response.status_code = 503
			return response
	except Exception as e:
		log('\n\n\n'+str(e)+'\n\n\n')
		return JsonResponse({'success':False, 'message':'Could not add figures to dashboard.'})
